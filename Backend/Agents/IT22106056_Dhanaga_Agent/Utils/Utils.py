import re
import os
import uuid
import json
import sqlite3
from datetime import datetime
from typing import Dict, List, Optional, Tuple
import pdfplumber
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import requests
import io

from .text_cleaners import TextCleaners
from .numeric_normalizers import NumericNormalizers
from .structure_cleaners import StructureCleaners

class DocumentProcessor:
    """
    Comprehensive document processor for climate policy documents
    Handles text extraction, cleaning, and preprocessing
    """

    def __init__(self, db_path: str = "processed_documents.db"):
        self.db_path = db_path
        self.init_database()
        self.text_cleaners = TextCleaners()
        self.numeric_normalizers = NumericNormalizers()
        self.structure_cleaners = StructureCleaners()

    def init_database(self):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS processed_documents (
                id TEXT PRIMARY KEY,
                filename TEXT NOT NULL,
                original_text TEXT NOT NULL,
                processed_text TEXT NOT NULL,
                metadata TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        conn.commit()
        conn.close()

    def extract_text(self, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            return self._extract_pdf_text(file_path)
        elif ext in ['.docx', '.doc']:
            return self._extract_docx_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def _extract_pdf_text(self, file_path: str) -> str:
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            # fallback with fitz
            try:
                doc = fitz.open(file_path)
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
            except Exception as e2:
                raise Exception(f"Both PDF extraction methods failed: pdfplumber: {e}, fitz(PyMuPDF): {e2}")
        return text

    def _extract_docx_text(self, file_path: str) -> str:
        doc = DocxDocument(file_path)
        text = ""
        for p in doc.paragraphs:
            text += p.text + "\n"
        return text

    def clean_and_preprocess(self, text: str) -> str:
        text = self.text_cleaners.basic_text_cleaning(text)
        text = self.text_cleaners.remove_headers_footers(text)
        text = self.text_cleaners.clean_special_characters(text)
        text = self.numeric_normalizers.normalize_numeric_and_dates(text)
        text = self.numeric_normalizers.handle_missing_values(text)
        text = self.structure_cleaners.structure_oriented_cleaning(text)
        text = self.structure_cleaners.remove_repeated_paragraphs(text)
        text = self.structure_cleaners.climate_policy_specific_cleaning(text)
        text = self.structure_cleaners.final_cleanup(text)
        return text

    def build_docx_bytes(self, text: str, title: str = "Processed Document"):
        doc = DocxDocument()
        doc.add_heading(title, level=1)
        for para in [p.strip() for p in text.split("\n\n") if p.strip()]:
            doc.add_paragraph(para)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def store_processed_document(self, filename: str, original_text: str, processed_text: str) -> str:
        doc_id = str(uuid.uuid4())
        metadata = {
            "original_length": len(original_text),
            "processed_length": len(processed_text),
            "processing_timestamp": datetime.now().isoformat(),
            "compression_ratio": len(processed_text) / len(original_text) if len(original_text) else 0,
        }
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO processed_documents (id, filename, original_text, processed_text, metadata)
            VALUES (?, ?, ?, ?, ?)
        ''', (doc_id, filename, original_text, processed_text, json.dumps(metadata)))
        conn.commit()
        conn.close()
        return doc_id

    def get_processed_document(self, doc_id: str) -> Optional[Dict]:
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            SELECT filename, original_text, processed_text, metadata, created_at
            FROM processed_documents WHERE id = ?
        ''', (doc_id,))
        row = cursor.fetchone()
        conn.close()
        if not row:
            return None
        return {
            "document_id": doc_id,
            "filename": row[0],
            "original_text": row[1],
            "processed_text": row[2],
            "metadata": json.loads(row[3]) if row[3] else {},
            "created_at": row[4],
        }

    def list_documents(self) -> List[Dict]:
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            SELECT id, filename, metadata, created_at
            FROM processed_documents ORDER BY created_at DESC
        ''')
        rows = cursor.fetchall()
        conn.close()
        return [
            {
                "document_id": r[0],
                "filename": r[1],
                "metadata": json.loads(r[2]) if r[2] else {},
                "created_at": r[3],
            }
            for r in rows
        ]

    def send_to_analyzer_agent(self, doc_id: str) -> Dict:
        """
        Send processed document to Document Analyzer Agent.
        Config via env vars:
        - ANALYZER_URL: target endpoint (required to actually send)
        - ANALYZER_SEND_FORMAT: 'json' (default) or 'docx'
        - ANALYZER_API_KEY: optional bearer token
        """
        document = self.get_processed_document(doc_id)
        if not document:
            raise ValueError(f"Document {doc_id} not found")

        analyzer_url = os.environ.get("ANALYZER_URL", "").strip()
        send_format = (os.environ.get("ANALYZER_SEND_FORMAT", "json").strip().lower() or "json")
        api_key = os.environ.get("ANALYZER_API_KEY", "").strip()

        meta = {
            "document_id": doc_id,
            "filename": document["filename"],
            "metadata": document.get("metadata", {}),
        }

        if not analyzer_url:
            return {
                "status": "mock",
                "message": "ANALYZER_URL not configured. Skipping external call.",
                "would_send_format": send_format,
                "payload_preview": {
                    "document_id": meta["document_id"],
                    "filename": meta["filename"],
                    "text_size": len(document["processed_text"]) if send_format == "json" else None
                }
            }

        headers = {"Accept": "application/json"}
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"

        try:
            if send_format == "docx":
                buf = self.build_docx_bytes(document["processed_text"], title=os.path.splitext(document["filename"])[0])
                files = {
                    "file": (
                        f"processed_{os.path.splitext(document['filename'])[0]}.docx",
                        buf,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                }
                data = {"metadata": json.dumps(meta)}
                resp = requests.post(analyzer_url, headers=headers, files=files, data=data, timeout=60)
            else:
                payload = {**meta, "processed_text": document["processed_text"]}
                headers["Content-Type"] = "application/json"
                resp = requests.post(analyzer_url, headers=headers, json=payload, timeout=60)

            try:
                content = resp.json()
            except Exception:
                content = {"text": resp.text}

            return {
                "status": "ok" if resp.ok else "error",
                "http_status": resp.status_code,
                "response": content,
            }
        except requests.RequestException as e:
            return {
                "status": "error",
                "message": f"Failed to send to analyzer agent: {e}",
            }