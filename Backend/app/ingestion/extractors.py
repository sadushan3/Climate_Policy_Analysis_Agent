"""Text extraction that preserves page provenance.

The V1 pipeline flattened a document to one string, which made citation
impossible -- you could not say *where* a claim came from. Here every extracted
block keeps its page number and its character span in the reconstructed
document, so any downstream claim can be traced back to a page.
"""
from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

from app.core.errors import NoTextExtracted, UnsupportedFileType

log = logging.getLogger(__name__)


@dataclass
class Page:
    number: int          # 1-indexed, as a human would cite it
    text: str
    char_start: int      # span within ExtractedDocument.text
    char_end: int


@dataclass
class ExtractedDocument:
    filename: str
    text: str
    pages: list[Page] = field(default_factory=list)
    extractor: str = ""
    meta: dict = field(default_factory=dict)

    def page_for_offset(self, offset: int) -> int:
        """Map a character offset in `text` back to a page number."""
        for page in self.pages:
            if page.char_start <= offset < page.char_end:
                return page.number
        return self.pages[-1].number if self.pages else 1


# Headers/footers repeat verbatim on most pages of a policy PDF and pollute both
# embeddings and keyword counts. Detect them by cross-page frequency rather than
# by position, which is robust to varying page layouts.
def _strip_repeated_lines(page_texts: list[str], min_ratio: float = 0.6) -> list[str]:
    if len(page_texts) < 4:
        return page_texts

    from collections import Counter

    counts: Counter[str] = Counter()
    for text in page_texts:
        lines = {ln.strip() for ln in text.splitlines() if 0 < len(ln.strip()) <= 90}
        counts.update(lines)

    threshold = len(page_texts) * min_ratio
    boilerplate = {
        line
        for line, count in counts.items()
        # A bare page number differs per page, so catch it by shape too.
        if count >= threshold and not re.fullmatch(r"[\d\W]{1,4}", line)
    }
    if boilerplate:
        log.debug("Removing %d boilerplate lines", len(boilerplate))

    cleaned = []
    for text in page_texts:
        kept = [ln for ln in text.splitlines() if ln.strip() not in boilerplate]
        cleaned.append("\n".join(kept))
    return cleaned


def _assemble(filename: str, page_texts: list[str], extractor: str, meta: dict) -> ExtractedDocument:
    page_texts = _strip_repeated_lines(page_texts)

    parts: list[str] = []
    pages: list[Page] = []
    cursor = 0
    for i, raw in enumerate(page_texts, start=1):
        normalised = normalise_whitespace(raw)
        if not normalised:
            continue
        block = normalised + "\n\n"
        pages.append(Page(number=i, text=normalised, char_start=cursor, char_end=cursor + len(block)))
        parts.append(block)
        cursor += len(block)

    text = "".join(parts).strip()
    if len(text) < 40:
        raise NoTextExtracted(
            "No usable text could be extracted. If this is a scanned document, "
            "it needs OCR before analysis."
        )
    return ExtractedDocument(filename=filename, text=text, pages=pages, extractor=extractor, meta=meta)


def normalise_whitespace(text: str) -> str:
    """Repair PDF text without destroying it.

    Deliberately preserves `%`, `$`, `°`, `€` and digits: in a climate policy the
    numbers *are* the content. V1's sanitiser stripped them, which silently
    deleted every emission target and finance commitment from the analysis.
    """
    # De-hyphenate words split across a line break ("emis-\nsions" -> "emissions").
    text = re.sub(r"(\w)-\s*\n\s*(\w)", r"\1\2", text)
    # Join lines that are a wrapped continuation of a sentence.
    text = re.sub(r"(?<![.!?:;])\n(?![\n\s•\-\d])", " ", text)
    # Collapse runs of blank lines to a single paragraph break.
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t\u00a0]+", " ", text)
    # Normalise unicode punctuation that breaks naive tokenisers.
    text = text.translate(str.maketrans({"\u2018": "'", "\u2019": "'", "\u201c": '"', "\u201d": '"', "\u2013": "-", "\u2014": "-"}))
    return "\n".join(line.strip() for line in text.splitlines()).strip()


def extract_pdf(data: bytes, filename: str) -> ExtractedDocument:
    import fitz  # PyMuPDF

    with fitz.open(stream=data, filetype="pdf") as doc:
        # "blocks" sorts by reading order, which handles two-column policy
        # layouts far better than the flat "text" mode used in V1.
        page_texts = [
            "\n".join(b[4] for b in page.get_text("blocks", sort=True) if b[4].strip())
            for page in doc
        ]
        meta = {"page_count": doc.page_count, **{k: v for k, v in (doc.metadata or {}).items() if v}}
    return _assemble(filename, page_texts, "pymupdf", meta)


def extract_docx(data: bytes, filename: str) -> ExtractedDocument:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    # DOCX has no reliable page model; page-break every ~3000 chars so citations
    # still point at a stable, roughly page-sized region.
    page_texts, buf, size = [], [], 0
    for part in parts:
        buf.append(part)
        size += len(part)
        if size > 3000:
            page_texts.append("\n".join(buf))
            buf, size = [], 0
    if buf:
        page_texts.append("\n".join(buf))

    return _assemble(filename, page_texts, "python-docx", {"synthetic_pages": True})


def extract_text(data: bytes, filename: str) -> ExtractedDocument:
    raw = data.decode("utf-8", errors="replace")
    chunks = [raw[i : i + 3000] for i in range(0, len(raw), 3000)] or [""]
    return _assemble(filename, chunks, "plaintext", {"synthetic_pages": True})


_EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".doc": extract_docx,
    ".txt": extract_text,
    ".md": extract_text,
}


def extract(data: bytes, filename: str) -> ExtractedDocument:
    suffix = Path(filename).suffix.lower()
    handler = _EXTRACTORS.get(suffix)
    if handler is None:
        raise UnsupportedFileType(
            f"'{suffix or filename}' is not a supported format.",
            details={"supported": sorted(_EXTRACTORS)},
        )
    return handler(data, filename)
